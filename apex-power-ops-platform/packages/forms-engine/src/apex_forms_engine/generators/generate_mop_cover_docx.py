"""
MOP Cover Page — Word (.docx) Template Generator
=================================================
Generates an editable Word document that mirrors the HTML cover page design.
Uses tables for layout since Word doesn't support CSS Grid/Flexbox.

Usage:
    python Scripts/generate_mop_cover_docx.py

Output:
    Templates/MOP/MOP_Cover_Page_Template.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Brand Colors ──
RESA_BLUE = RGBColor(0x01, 0x56, 0x87)
RESA_GREEN = RGBColor(0x5F, 0xA8, 0x44)
RESA_GRAY = RGBColor(0x68, 0x68, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
BORDER_GRAY = RGBColor(0xDD, 0xDD, 0xDD)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)

# ── Paths ──
PACKAGE_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
NETA_LOGO = os.path.join(PACKAGE_ROOT, "assets", "neta-logo-90px.png")
RESA_LOGO = os.path.join(PACKAGE_ROOT, "assets", "resa-logo-90px.png")
OUTPUT_PATH = os.path.join(PACKAGE_ROOT, "artifacts", "MOP_Cover_Page_Template.docx")


# ── Helper Functions ──

def set_cell_shading(cell, hex_color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. Usage: set_cell_border(cell, top={"sz": 4, "color": "015687"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
            tcBorders.append(element)
        for key, val in attrs.items():
            element.set(qn(f"w:{key}"), str(val))


def add_formatted_run(paragraph, text, size=8, bold=False, color=BLACK, italic=False, caps=False):
    """Add a run with formatting."""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    if caps:
        run.font.all_caps = True
    run.font.name = "Segoe UI"
    return run


def set_row_height(row, height_pt):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="exact"/>')
    trPr.append(trHeight)


def set_cell_width(cell, width_inches):
    """Set cell width."""
    cell.width = Inches(width_inches)


def make_section_header(doc, text):
    """Create a RESA-blue section header bar."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "015687")
    p = cell.paragraphs[0]
    add_formatted_run(p, text, size=8, bold=True, color=WHITE, caps=True)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Remove table borders (the shading IS the visual bar)
    for border_name in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "015687"}})
    return table


def add_spacer(doc, pts=4):
    """Add a small spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    pf = p.paragraph_format
    pf.line_spacing = Pt(pts)
    return p


# ══════════════════════════════════════════════
#  MAIN DOCUMENT BUILD
# ══════════════════════════════════════════════

def build_document():
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Segoe UI"
    font.size = Pt(8)
    font.color.rgb = BLACK

    usable_width = 7.5  # 8.5 - 0.5 - 0.5

    # ═══════════════════  HEADER  ═══════════════════
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(3.5)
    header_table.columns[2].width = Inches(2.0)

    # NETA logo (left)
    cell_left = header_table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = cell_left.paragraphs[0]
    if os.path.exists(NETA_LOGO):
        p.add_run().add_picture(NETA_LOGO, height=Inches(0.5))
    else:
        add_formatted_run(p, "[NETA Logo]", size=7, color=RESA_GRAY, italic=True)

    # Center cell - empty
    center = header_table.cell(0, 1)
    center.paragraphs[0].text = ""

    # RESA logo (right)
    cell_right = header_table.cell(0, 2)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(RESA_LOGO):
        p.add_run().add_picture(RESA_LOGO, height=Inches(0.5))
    else:
        add_formatted_run(p, "[RESA Logo]", size=7, color=RESA_GRAY, italic=True)

    # Remove header table borders
    for row in header_table.rows:
        for cell in row.cells:
            for border_name in ["top", "left", "bottom", "right"]:
                set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "FFFFFF"}})

    # ── Blue line ──
    blue_line = doc.add_table(rows=1, cols=1)
    blue_line.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = blue_line.cell(0, 0)
    set_cell_shading(cell, "015687")
    set_row_height(blue_line.rows[0], 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for border_name in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "015687"}})

    # ── Tagline + Address row ──
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.columns[0].width = Inches(4.0)
    info_table.columns[1].width = Inches(3.5)

    tagline_cell = info_table.cell(0, 0)
    p = tagline_cell.paragraphs[0]
    add_formatted_run(p, "Reliable and Safe. ", size=7.5, bold=True, color=RESA_BLUE, italic=True)
    add_formatted_run(p, "The Power of Us.", size=7.5, bold=True, color=RESA_GREEN, italic=True)

    addr_cell = info_table.cell(0, 1)
    p = addr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, "2447 W 12th St. Suite 5 · Tempe, AZ 85142", size=7.5, color=RESA_GRAY)

    for row in info_table.rows:
        for cell in row.cells:
            for border_name in ["top", "left", "bottom", "right"]:
                set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "FFFFFF"}})

    # ═══════════════════  DOCUMENT TITLE  ═══════════════════
    add_spacer(doc, 4)
    p = doc.add_paragraph()
    run = add_formatted_run(p, "METHOD OF PROCEDURE", size=13, bold=True, color=RESA_BLUE)
    # Green underline using bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="2" w:color="5FA844"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ═══════════════════  DOCUMENT CONTROL BAR  ═══════════════════
    add_spacer(doc, 2)
    ctrl_table = doc.add_table(rows=2, cols=4)
    ctrl_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels = ["MOP ID", "VERSION", "STATUS", "DATE"]
    placeholders = ["{{mop_id}}", "{{version}}", "{{status}}", "{{date}}"]

    for i in range(4):
        ctrl_table.columns[i].width = Inches(usable_width / 4)
        # Label row
        cell = ctrl_table.cell(0, i)
        set_cell_shading(cell, "015687")
        p = cell.paragraphs[0]
        add_formatted_run(p, labels[i], size=6.5, color=WHITE, caps=True)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        for border_name in ["top", "left", "bottom", "right"]:
            set_cell_border(cell, **{border_name: {"sz": 2, "val": "single", "color": "FFFFFF"}})

        # Value row
        cell = ctrl_table.cell(1, i)
        set_cell_shading(cell, "015687")
        p = cell.paragraphs[0]
        add_formatted_run(p, placeholders[i], size=8, bold=True, color=WHITE)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        for border_name in ["top", "left", "bottom", "right"]:
            set_cell_border(cell, **{border_name: {"sz": 2, "val": "single", "color": "FFFFFF"}})

    # ═══════════════════  PROJECT INFORMATION  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Project Information")

    proj_table = doc.add_table(rows=4, cols=4)
    proj_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    proj_fields = [
        ("PROJECT:", "{{project_name}}", "JOB NUMBER:", "{{job_number}}"),
        ("CLIENT:", "{{client_name}}", "OWNER:", "{{owner_name}}"),
        ("SITE:", "{{site_name}}", "ADDRESS:", "{{site_address}}"),
        ("START / END:", "{{start_datetime}} — {{end_datetime}}", "CLASSIFICATION:", "{{classification}}"),
    ]

    for row_idx, (lbl1, val1, lbl2, val2) in enumerate(proj_fields):
        row = proj_table.rows[row_idx]
        # Column widths
        proj_table.columns[0].width = Inches(1.0)
        proj_table.columns[1].width = Inches(2.75)
        proj_table.columns[2].width = Inches(1.0)
        proj_table.columns[3].width = Inches(2.75)

        for col_idx, (label, value) in enumerate([(lbl1, val1), (lbl2, val2)]):
            label_cell = row.cells[col_idx * 2]
            value_cell = row.cells[col_idx * 2 + 1]

            p = label_cell.paragraphs[0]
            add_formatted_run(p, label, size=6.5, bold=True, color=RESA_GRAY, caps=True)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            p = value_cell.paragraphs[0]
            add_formatted_run(p, value, size=8, color=BLACK)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            # Borders
            for cell in [label_cell, value_cell]:
                set_cell_border(cell,
                    top={"sz": 0, "val": "none", "color": "FFFFFF"},
                    left={"sz": 0, "val": "none", "color": "FFFFFF"},
                    right={"sz": 0, "val": "none", "color": "FFFFFF"},
                    bottom={"sz": 4, "val": "single", "color": "015687"},
                )

    # outer border for section
    tbl = proj_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # ═══════════════════  APPROVAL TABLE  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "MOP Review — Approval to Proceed")

    approval_table = doc.add_table(rows=5, cols=4)
    approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    approval_table.columns[0].width = Inches(2.1)
    approval_table.columns[1].width = Inches(1.875)
    approval_table.columns[2].width = Inches(2.025)
    approval_table.columns[3].width = Inches(1.5)

    # Header row
    headers = ["ROLE / AUTHORITY", "NAME", "SIGNATURE", "DATE"]
    for i, h in enumerate(headers):
        cell = approval_table.cell(0, i)
        set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        add_formatted_run(p, h, size=6.5, bold=True, color=RESA_GRAY, caps=True)
        set_cell_border(cell,
            top={"sz": 4, "val": "single", "color": "DDDDDD"},
            left={"sz": 4, "val": "single", "color": "DDDDDD"},
            bottom={"sz": 4, "val": "single", "color": "DDDDDD"},
            right={"sz": 4, "val": "single", "color": "DDDDDD"},
        )

    # Data rows
    for row_idx in range(1, 5):
        for col_idx in range(4):
            cell = approval_table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            if col_idx == 0:
                add_formatted_run(p, f"{{{{approval_role_{row_idx}}}}}", size=7.5, color=BLACK)
            elif col_idx == 1:
                add_formatted_run(p, f"{{{{approval_name_{row_idx}}}}}", size=7.5, color=BLACK)
            else:
                add_formatted_run(p, "", size=7.5)  # blank for signature/date
            set_cell_border(cell,
                top={"sz": 4, "val": "single", "color": "DDDDDD"},
                left={"sz": 4, "val": "single", "color": "DDDDDD"},
                bottom={"sz": 4, "val": "single", "color": "DDDDDD"},
                right={"sz": 4, "val": "single", "color": "DDDDDD"},
            )
            set_row_height(approval_table.rows[row_idx], 18)

    # ═══════════════════  ATTENDEES TABLE  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Attendees — Roles & Responsibilities")

    att_table = doc.add_table(rows=6, cols=4)
    att_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    att_table.columns[0].width = Inches(1.875)
    att_table.columns[1].width = Inches(1.875)
    att_table.columns[2].width = Inches(1.875)
    att_table.columns[3].width = Inches(1.875)

    att_headers = ["ROLE", "NAME", "COMPANY", "PHONE"]
    for i, h in enumerate(att_headers):
        cell = att_table.cell(0, i)
        set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        add_formatted_run(p, h, size=6.5, bold=True, color=RESA_GRAY, caps=True)
        set_cell_border(cell,
            top={"sz": 4, "val": "single", "color": "DDDDDD"},
            left={"sz": 4, "val": "single", "color": "DDDDDD"},
            bottom={"sz": 4, "val": "single", "color": "DDDDDD"},
            right={"sz": 4, "val": "single", "color": "DDDDDD"},
        )

    att_fields = ["role", "name", "company", "phone"]
    for row_idx in range(1, 6):
        for col_idx, field in enumerate(att_fields):
            cell = att_table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            add_formatted_run(p, f"{{{{attendee_{field}_{row_idx}}}}}", size=7.5, color=BLACK)
            set_cell_border(cell,
                top={"sz": 4, "val": "single", "color": "DDDDDD"},
                left={"sz": 4, "val": "single", "color": "DDDDDD"},
                bottom={"sz": 4, "val": "single", "color": "DDDDDD"},
                right={"sz": 4, "val": "single", "color": "DDDDDD"},
            )

    # ═══════════════════  SUMMARY OF WORK  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Summary of Work — Facility Impact")

    # Impact box (yellow left-border style simulated with a table)
    impact_table = doc.add_table(rows=5, cols=2)
    impact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    impact_table.columns[0].width = Inches(1.5)
    impact_table.columns[1].width = Inches(6.0)

    impact_fields = [
        ("FACILITY IMPACT:", "{{facility_impact}}"),
        ("SYSTEMS AFFECTED:", "{{systems_affected}}"),
        ("OPERATIONAL RISK:", "{{operational_risk}}"),
        ("EST. DURATION:", "{{estimated_duration}}"),
        ("OUTAGE WINDOW:", "{{outage_window}}"),
    ]

    for row_idx, (label, value) in enumerate(impact_fields):
        label_cell = impact_table.cell(row_idx, 0)
        value_cell = impact_table.cell(row_idx, 1)

        p = label_cell.paragraphs[0]
        add_formatted_run(p, label, size=7, bold=True, color=RESA_GRAY, caps=True)

        p = value_cell.paragraphs[0]
        add_formatted_run(p, value, size=8, color=BLACK)

        # Yellow left border on label cells
        set_cell_border(label_cell,
            left={"sz": 12, "val": "single", "color": "FFC107"},
            top={"sz": 0, "val": "none", "color": "FFFFFF"},
            bottom={"sz": 0, "val": "none", "color": "FFFFFF"},
            right={"sz": 0, "val": "none", "color": "FFFFFF"},
        )
        set_cell_border(value_cell,
            top={"sz": 0, "val": "none", "color": "FFFFFF"},
            bottom={"sz": 0, "val": "none", "color": "FFFFFF"},
            left={"sz": 0, "val": "none", "color": "FFFFFF"},
            right={"sz": 0, "val": "none", "color": "FFFFFF"},
        )
        set_cell_shading(label_cell, "FFF8E6")
        set_cell_shading(value_cell, "FFF8E6")

    # Scope & Work Includes
    add_spacer(doc, 2)
    p = doc.add_paragraph()
    add_formatted_run(p, "Scope: ", size=7.5, bold=True, color=RESA_BLUE)
    add_formatted_run(p, "{{scope_description}}", size=7.5, color=BLACK)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    add_formatted_run(p, "Work Includes: ", size=7.5, bold=True, color=RESA_BLUE)
    add_formatted_run(p, "{{work_includes}}", size=7.5, color=BLACK)
    p.paragraph_format.space_after = Pt(2)

    # ═══════════════════  SAFETY REQUIREMENTS  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Safety Requirements")

    safety_table = doc.add_table(rows=5, cols=6)
    safety_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column headers
    safety_cols = [
        ("ARC FLASH DETAILS", [
            ("Arc Flash Boundary:", "{{arc_flash_boundary}}"),
            ("Incident Energy:", "{{incident_energy}}"),
            ("PPE Category:", "{{ppe_category}}"),
            ("Working Distance:", "{{working_distance}}"),
        ]),
        ("EMERGENCY INFORMATION", [
            ("Emergency Contact:", "{{emergency_contact}}"),
            ("Emergency Phone:", "{{emergency_phone}}"),
            ("Assembly Point:", "{{assembly_point}}"),
            ("", ""),
        ]),
        ("MEDICAL RESPONSE", [
            ("Nearest Hospital:", "{{nearest_hospital}}"),
            ("AED Location:", "{{aed_location}}"),
            ("First Aid Kit:", "{{first_aid_location}}"),
            ("", ""),
        ]),
    ]

    for col_group_idx, (title, fields) in enumerate(safety_cols):
        label_col = col_group_idx * 2
        value_col = col_group_idx * 2 + 1

        # Title row (merged)
        title_cell = safety_table.cell(0, label_col)
        title_cell_merge = safety_table.cell(0, value_col)
        title_cell.merge(title_cell_merge)
        p = title_cell.paragraphs[0]
        add_formatted_run(p, title, size=7, bold=True, color=RESA_BLUE, caps=True)
        # Green underline
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        for row_idx, (label, value) in enumerate(fields):
            r = row_idx + 1
            cell_l = safety_table.cell(r, label_col)
            cell_v = safety_table.cell(r, value_col)

            p = cell_l.paragraphs[0]
            add_formatted_run(p, label, size=7, bold=True, color=RESA_GRAY)

            p = cell_v.paragraphs[0]
            add_formatted_run(p, value, size=7, color=BLACK)

    # Remove all borders from safety table
    for row in safety_table.rows:
        for cell in row.cells:
            for border_name in ["top", "left", "bottom", "right"]:
                set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "FFFFFF"}})

    # Outer border for section
    tbl = safety_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # ── Safety statement ──
    add_spacer(doc, 2)
    p = doc.add_paragraph()
    add_formatted_run(p, "RESA Power Service", size=7, bold=True, color=RESA_BLUE)
    add_formatted_run(p, " is committed to the safety of all personnel involved in this work. ", size=7, color=BLACK)
    add_formatted_run(p, "All activities will be performed in strict adherence to ", size=7, color=BLACK)
    add_formatted_run(p, "NFPA 70E", size=7, bold=True, color=RESA_BLUE)
    add_formatted_run(p, " and ", size=7, color=BLACK)
    add_formatted_run(p, "OSHA 29 CFR 1910", size=7, bold=True, color=RESA_BLUE)
    add_formatted_run(p, " regulations.", size=7, color=BLACK)

    # ── Stop Work Authority ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "⚠ STOP WORK AUTHORITY: ", size=7, bold=True, color=RGBColor(0x72, 0x1C, 0x24))
    add_formatted_run(p, "All personnel have the authority and responsibility to stop work immediately if unsafe conditions are observed.", size=7, color=RGBColor(0x72, 0x1C, 0x24))

    # ═══════════════════  GENERAL NOTES  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "General Notes")

    notes_table = doc.add_table(rows=1, cols=2)
    notes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    notes_table.columns[0].width = Inches(3.75)
    notes_table.columns[1].width = Inches(3.75)

    # Left column - MOP Structure
    left_cell = notes_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    add_formatted_run(p, "MOP STRUCTURE & SECTIONS", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    sections = [
        ("Sec 2", "Scope Applicability & Requirements"),
        ("Sec 3", "Job Hazard Analysis (JHA)"),
        ("Sec 4", "Required Forms & Documentation"),
        ("Sec 5", "Work Script & Hold Points"),
        ("Sec 6", "Completion & Sign-Off"),
    ]
    for badge, desc in sections:
        p = left_cell.add_paragraph()
        add_formatted_run(p, f"  [{badge}]  ", size=6, bold=True, color=WHITE)
        # Can't truly shade inline text, so use brackets
        add_formatted_run(p, desc, size=7, color=BLACK)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    # Right column - Key Information
    right_cell = notes_table.cell(0, 1)
    p = right_cell.paragraphs[0]
    add_formatted_run(p, "KEY INFORMATION", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    key_notes = [
        "All personnel must review this MOP before work begins",
        "Changes after approval require re-authorization",
        "Hold Points require verification before proceeding",
        "Document all deviations in Execution Summary",
    ]
    for note in key_notes:
        p = right_cell.add_paragraph()
        add_formatted_run(p, f"• {note}", size=7, color=BLACK)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    # Remove borders
    for row in notes_table.rows:
        for cell in row.cells:
            for border_name in ["top", "left", "bottom", "right"]:
                set_cell_border(cell, **{border_name: {"sz": 4, "val": "single", "color": "015687"}})

    # ═══════════════════  FOOTER  ═══════════════════
    add_spacer(doc, 4)
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.columns[0].width = Inches(1.5)
    footer_table.columns[1].width = Inches(6.0)

    green_cell = footer_table.cell(0, 0)
    set_cell_shading(green_cell, "5FA844")
    p = green_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "resapower.com", size=6, bold=True, color=WHITE)

    blue_cell = footer_table.cell(0, 1)
    set_cell_shading(blue_cell, "015687")
    p = blue_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, "RESA Power Service  ·  NETA Accredited Company  ·  Page {{page_number}} of {{total_pages}}", size=6, color=WHITE)

    for cell in [green_cell, blue_cell]:
        for border_name in ["top", "left", "bottom", "right"]:
            set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "FFFFFF"}})

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"✅ Word template saved: {OUTPUT_PATH}")
    print(f"   File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
