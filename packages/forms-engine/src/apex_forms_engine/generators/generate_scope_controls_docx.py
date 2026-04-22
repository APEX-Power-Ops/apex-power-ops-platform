"""
MOP Page 2 — Scope & Controls Word (.docx) Template Generator
==============================================================
Generates an editable Word document matching the HTML Scope & Controls page.
Includes: Scope Applicability, Pre-Energization Checklist, Safety Controls,
Hold Points & Authority, Required Forms & Referenced Documents.

Usage:
    python Scripts/generate_scope_controls_docx.py

Output:
    Templates/MOP/MOP_Scope_Controls_Template.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Brand Colors ──
RESA_BLUE = RGBColor(0x01, 0x56, 0x87)
RESA_GREEN = RGBColor(0x5F, 0xA8, 0x44)
RESA_GRAY = RGBColor(0x68, 0x68, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)
RISK_HIGH_RGB = RGBColor(0xDC, 0x35, 0x45)

# ── Paths ──
PACKAGE_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
NETA_LOGO = os.path.join(PACKAGE_ROOT, "assets", "neta-logo-90px.png")
RESA_LOGO = os.path.join(PACKAGE_ROOT, "assets", "resa-logo-90px.png")
OUTPUT_PATH = os.path.join(PACKAGE_ROOT, "artifacts", "MOP_Scope_Controls_Template.docx")


# ── Helper Functions ──

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
            tcBorders.append(element)
        for key, val in attrs.items():
            element.set(qn(f"w:{key}"), str(val))


def add_formatted_run(paragraph, text, size=8, bold=False, color=BLACK, italic=False, caps=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    if caps:
        run.font.all_caps = True
    run.font.name = "Segoe UI"
    return run


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="exact"/>')
    trPr.append(trHeight)


def make_section_header(doc, text, alert=False):
    """Create a RESA-blue (or red alert) section header bar."""
    color_hex = "DC3545" if alert else "015687"
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color_hex)
    p = cell.paragraphs[0]
    add_formatted_run(p, text, size=8, bold=True, color=WHITE, caps=True)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for border_name in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": color_hex}})
    return table


def add_spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.line_spacing = Pt(pts)
    return p


def remove_all_borders(cell):
    for border_name in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "FFFFFF"}})


def add_branded_header(doc):
    """Add the standard RESA branded header with logos, blue line, tagline, address."""
    # Logo row
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(3.5)
    header_table.columns[2].width = Inches(2.0)

    cell_left = header_table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = cell_left.paragraphs[0]
    if os.path.exists(NETA_LOGO):
        p.add_run().add_picture(NETA_LOGO, height=Inches(0.5))
    else:
        add_formatted_run(p, "[NETA Logo]", size=7, color=RESA_GRAY, italic=True)

    header_table.cell(0, 1).paragraphs[0].text = ""

    cell_right = header_table.cell(0, 2)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(RESA_LOGO):
        p.add_run().add_picture(RESA_LOGO, height=Inches(0.5))
    else:
        add_formatted_run(p, "[RESA Logo]", size=7, color=RESA_GRAY, italic=True)

    for row in header_table.rows:
        for cell in row.cells:
            remove_all_borders(cell)

    # Blue line
    blue_line = doc.add_table(rows=1, cols=1)
    blue_line.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = blue_line.cell(0, 0)
    set_cell_shading(cell, "015687")
    set_row_height(blue_line.rows[0], 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for b in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{b: {"sz": 0, "val": "none", "color": "015687"}})

    # Tagline + Address
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.columns[0].width = Inches(4.0)
    info_table.columns[1].width = Inches(3.5)

    p = info_table.cell(0, 0).paragraphs[0]
    add_formatted_run(p, "Reliable and Safe. ", size=7.5, bold=True, color=RESA_BLUE, italic=True)
    add_formatted_run(p, "The Power of Us.", size=7.5, bold=True, color=RESA_GREEN, italic=True)

    p = info_table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, "2447 W 12th St. Suite 5 · Tempe, AZ 85142", size=7.5, color=RESA_GRAY)

    for row in info_table.rows:
        for cell in row.cells:
            remove_all_borders(cell)


def add_doc_info_bar(doc):
    """Add the document info bar (MOP type, ID, project/site)."""
    bar = doc.add_table(rows=1, cols=2)
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar.columns[0].width = Inches(4.0)
    bar.columns[1].width = Inches(3.5)

    cell = bar.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    add_formatted_run(p, "Method of Procedure  ", size=9, bold=True, color=RESA_BLUE)
    add_formatted_run(p, "{{mop_id}} · v{{version}}", size=7, color=RESA_GRAY)

    cell = bar.cell(0, 1)
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, "{{project_name}}\n{{site_name}}", size=7, color=RESA_GRAY)

    for cell in bar.row_cells(0):
        set_cell_border(cell,
            top={"sz": 0, "val": "none", "color": "FFFFFF"},
            left={"sz": 0, "val": "none", "color": "FFFFFF"},
            right={"sz": 0, "val": "none", "color": "FFFFFF"},
            bottom={"sz": 4, "val": "single", "color": "DDDDDD"},
        )


def add_branded_footer(doc, page_text="Page 2 of {{total_pages}}"):
    """Add the green/blue footer bar."""
    add_spacer(doc, 4)
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.columns[0].width = Inches(1.5)
    footer_table.columns[1].width = Inches(6.0)

    green_cell = footer_table.cell(0, 0)
    set_cell_shading(green_cell, "5FA844")
    p = green_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "resapower.com", size=6, bold=True, color=WHITE)

    blue_cell = footer_table.cell(0, 1)
    set_cell_shading(blue_cell, "015687")
    p = blue_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, f"RESA Power Service  ·  NETA Accredited Company  ·  {page_text}", size=6, color=WHITE)

    for cell in [green_cell, blue_cell]:
        remove_all_borders(cell)


def add_checkbox_table(doc, title, items, col_width=3.75):
    """Add a column with a title and checkbox items."""
    # This is used inside a two-column layout, so we just return items for table cells
    pass  # Handled directly in build


# ══════════════════════════════════════════════
#  MAIN DOCUMENT BUILD
# ══════════════════════════════════════════════

def build_document():
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(8)
    style.font.color.rgb = BLACK

    # ═══════════════════  HEADER  ═══════════════════
    add_branded_header(doc)
    add_doc_info_bar(doc)
    add_spacer(doc, 4)

    # ═══════════════════  SCOPE + PRE-ENERGIZATION  ═══════════════════
    make_section_header(doc, "Scope Applicability & Pre-Energization Checklist")

    scope_items = [
        ("MV Switching", True), ("LV Switching", False), ("Relay Testing", True),
        ("Protection Testing", True), ("Commissioning", False), ("LOTO Required", True),
        ("CT Circuits", False), ("PT Circuits", False), ("Outage Coordination", False),
        ("Control Center Notify", True), ("Hot Work Permit", False), ("Confined Space", False),
    ]
    pre_energize_items = [
        "Pre-Functional checklist completed", "Vendor start-up performed (if applicable)",
        "Equipment ready for energization", "Proper voltage verified at source",
        "Feed verification per drawings", "Test equipment staged & calibrated",
        "Load bank connected (if applicable)", "Fire system in test mode",
        "Cooling/ventilation operational", "Barriers & signage in place",
        "All personnel briefed", "Communication plan confirmed",
    ]

    num_rows = max(len(scope_items), len(pre_energize_items))
    checklist_table = doc.add_table(rows=num_rows + 1, cols=4)
    checklist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    checklist_table.columns[0].width = Inches(0.35)
    checklist_table.columns[1].width = Inches(3.40)
    checklist_table.columns[2].width = Inches(0.35)
    checklist_table.columns[3].width = Inches(3.40)

    # Column titles
    cell = checklist_table.cell(0, 0)
    cell.merge(checklist_table.cell(0, 1))
    p = cell.paragraphs[0]
    add_formatted_run(p, "SCOPE APPLICABILITY", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/></w:pBdr>'))

    cell = checklist_table.cell(0, 2)
    cell.merge(checklist_table.cell(0, 3))
    p = cell.paragraphs[0]
    add_formatted_run(p, "PRE-ENERGIZATION CHECKLIST", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/></w:pBdr>'))

    # Checklist rows
    for i in range(num_rows):
        row = checklist_table.rows[i + 1]
        # Left: Scope
        if i < len(scope_items):
            label, checked = scope_items[i]
            p = row.cells[0].paragraphs[0]
            add_formatted_run(p, "☑" if checked else "☐", size=9, color=RESA_GREEN if checked else RESA_BLUE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = row.cells[1].paragraphs[0]
            add_formatted_run(p, label, size=7.5, color=BLACK)

        # Right: Pre-Energization
        if i < len(pre_energize_items):
            p = row.cells[2].paragraphs[0]
            add_formatted_run(p, "☐", size=9, color=RESA_BLUE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = row.cells[3].paragraphs[0]
            add_formatted_run(p, pre_energize_items[i], size=7.5, color=BLACK)

    # Borders for section
    for row in checklist_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 0, "val": "none", "color": "FFFFFF"},
                left={"sz": 0, "val": "none", "color": "FFFFFF"},
                right={"sz": 0, "val": "none", "color": "FFFFFF"},
                bottom={"sz": 0, "val": "none", "color": "FFFFFF"},
            )

    # Add outer border
    tbl = checklist_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'</w:tblBorders>'
    ))

    # ═══════════════════  SAFETY CONTROLS SUMMARY  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Safety Controls Summary")

    safety_data = [
        ("Voltage Class", "{{voltage_class}}"),
        ("Arc Flash Boundary", "{{arc_flash_boundary}}"),
        ("Incident Energy", "{{incident_energy}}"),
        ("PPE Category", "{{ppe_category}}"),
        ("Working Distance", "{{working_distance}}"),
        ("Limited Approach", "{{limited_approach}}"),
        ("Restricted Approach", "{{restricted_approach}}"),
        ("Prohibited Approach", "{{prohibited_approach}}"),
    ]

    safety_table = doc.add_table(rows=3, cols=6)
    safety_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(safety_data):
        row_idx = i // 3
        col_idx = (i % 3) * 2
        cell_l = safety_table.cell(row_idx, col_idx)
        p = cell_l.paragraphs[0]
        add_formatted_run(p, label, size=7, color=RESA_GRAY)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

        cell_v = safety_table.cell(row_idx, col_idx + 1)
        p = cell_v.paragraphs[0]
        add_formatted_run(p, value, size=7, bold=True, color=RESA_BLUE)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for row in safety_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 0, "val": "none", "color": "FFFFFF"},
                left={"sz": 0, "val": "none", "color": "FFFFFF"},
                right={"sz": 0, "val": "none", "color": "FFFFFF"},
                bottom={"sz": 2, "val": "single", "color": "F5F5F5"},
            )

    # Outer border
    tbl = safety_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'</w:tblBorders>'
    ))

    # Reference box
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="4" w:color="5FA844"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>'))
    add_formatted_run(p, "See AHA-{{aha_id}}", size=7, bold=True, color=RESA_BLUE)
    add_formatted_run(p, " for detailed hazard analysis, controls, and PPE requirements. ", size=7, color=BLACK)
    add_formatted_run(p, "All boundaries per NFPA 70E Table 130.4(D)(a).", size=7, italic=True, color=BLACK)

    # ═══════════════════  HOLD POINTS & AUTHORITY  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Hold Points & Authority")

    hp_table = doc.add_table(rows=5, cols=4)
    hp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hp_table.columns[0].width = Inches(0.6)
    hp_table.columns[1].width = Inches(3.15)
    hp_table.columns[2].width = Inches(1.5)
    hp_table.columns[3].width = Inches(2.25)

    # Left title
    cell = hp_table.cell(0, 0)
    cell.merge(hp_table.cell(0, 1))
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    add_formatted_run(p, "HOLD POINTS", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/></w:pBdr>'))

    # Right title
    cell = hp_table.cell(0, 2)
    cell.merge(hp_table.cell(0, 3))
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    add_formatted_run(p, "DECISION AUTHORITY", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/></w:pBdr>'))

    # Hold points data
    hold_points = [
        ("HP1", "Authorization to Proceed"),
        ("HP2", "Zero Voltage Verified"),
        ("HP3", "Pre-Energization Check"),
        ("HP4", "Final Sign-Off"),
    ]
    authority_data = [
        ("Stop Work", "All Personnel"),
        ("Back-Out Decision", "Crew Lead + Customer"),
        ("Scope Change", "PM Approval Required"),
        ("Est. Back-Out Time", "{{backout_time}}"),
    ]

    for i in range(4):
        row = hp_table.rows[i + 1]
        # HP badge
        p = row.cells[0].paragraphs[0]
        add_formatted_run(p, hold_points[i][0], size=7, bold=True, color=WHITE)
        set_cell_shading(row.cells[0], "DC3545")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # HP description
        p = row.cells[1].paragraphs[0]
        set_cell_shading(row.cells[1], "F5F5F5")
        add_formatted_run(p, hold_points[i][1], size=7.5, color=BLACK)

        # Authority label
        p = row.cells[2].paragraphs[0]
        set_cell_shading(row.cells[2], "F5F5F5")
        add_formatted_run(p, authority_data[i][0], size=7, color=RESA_GRAY)

        # Authority value
        p = row.cells[3].paragraphs[0]
        set_cell_shading(row.cells[3], "F5F5F5")
        add_formatted_run(p, authority_data[i][1], size=7, bold=True, color=BLACK)

    for row in hp_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 2, "val": "single", "color": "DDDDDD"},
                left={"sz": 2, "val": "single", "color": "DDDDDD"},
                right={"sz": 2, "val": "single", "color": "DDDDDD"},
                bottom={"sz": 2, "val": "single", "color": "DDDDDD"},
            )

    # Back-out reference
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="4" w:color="5FA844"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>'))
    add_formatted_run(p, "Back-Out Procedure: ", size=7, bold=True, color=RESA_BLUE)
    add_formatted_run(p, "See ", size=7, color=BLACK)
    add_formatted_run(p, "SOP-{{backout_sop}}", size=7, bold=True, color=RESA_BLUE)
    add_formatted_run(p, " for detailed back-out steps.", size=7, color=BLACK)

    # ═══════════════════  REQUIRED FORMS  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Required Forms & Referenced Documents")

    forms_data = [
        ("☑", "IVF-001", "Isolation Verification Form",     "☑", "PEC-001",  "Pre-Energization Checklist"),
        ("☑", "LCF-001", "Lifted Conductor Form",           "☐", "RSR-001",  "Relay Settings Record"),
        ("☑", "FTR-001", "Functional Test Record",           "☐", "POEC-001", "Post-Energization Checklist"),
    ]

    forms_table = doc.add_table(rows=len(forms_data) + 1, cols=6)
    forms_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    forms_table.columns[0].width = Inches(0.35)
    forms_table.columns[1].width = Inches(0.8)
    forms_table.columns[2].width = Inches(2.60)
    forms_table.columns[3].width = Inches(0.35)
    forms_table.columns[4].width = Inches(0.8)
    forms_table.columns[5].width = Inches(2.60)

    # Header row
    headers = ["Req", "Form ID", "Description", "Req", "Form ID", "Description"]
    for i, h in enumerate(headers):
        cell = forms_table.cell(0, i)
        set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        add_formatted_run(p, h, size=6.5, bold=True, color=RESA_GRAY, caps=True)
        if i in [0, 3]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(forms_data):
        row = forms_table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            p = row.cells[col_idx].paragraphs[0]
            if col_idx in [0, 3]:
                add_formatted_run(p, val, size=9, color=RESA_GREEN if val == "☑" else RESA_BLUE)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col_idx in [1, 4]:
                add_formatted_run(p, val, size=7, bold=True, color=RESA_BLUE)
            else:
                add_formatted_run(p, val, size=7, color=BLACK)

    for row in forms_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 2, "val": "single", "color": "DDDDDD"},
                left={"sz": 2, "val": "single", "color": "DDDDDD"},
                right={"sz": 2, "val": "single", "color": "DDDDDD"},
                bottom={"sz": 2, "val": "single", "color": "DDDDDD"},
            )

    # References
    add_spacer(doc, 2)
    ref_table = doc.add_table(rows=3, cols=2)
    ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ref_table.columns[0].width = Inches(0.75)
    ref_table.columns[1].width = Inches(6.75)

    refs = [
        ("AHA", "AHA-{{aha_id}} — Activity Hazard Analysis"),
        ("SOP", "{{sop_refs}}"),
        ("Drawings", "{{drawing_refs}}"),
    ]
    for i, (type_label, docs) in enumerate(refs):
        p = ref_table.cell(i, 0).paragraphs[0]
        add_formatted_run(p, type_label, size=6.5, bold=True, color=RESA_GRAY, caps=True)
        p = ref_table.cell(i, 1).paragraphs[0]
        add_formatted_run(p, docs, size=7, color=RESA_BLUE)

    for row in ref_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 2, "val": "single", "color": "DDDDDD"},
                left={"sz": 0, "val": "none", "color": "FFFFFF"},
                right={"sz": 0, "val": "none", "color": "FFFFFF"},
                bottom={"sz": 0, "val": "none", "color": "FFFFFF"},
            )

    # ═══════════════════  FOOTER  ═══════════════════
    add_branded_footer(doc, "Page 2 of {{total_pages}}")

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"✅ Word template saved: {OUTPUT_PATH}")
    print(f"   File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
