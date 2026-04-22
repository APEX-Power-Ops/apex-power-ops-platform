"""
MOP Final Page — Completion & Approvals Word (.docx) Template Generator
========================================================================
Generates an editable Word document matching the HTML Completion page.
Includes: Post-Work Verification, Restoration Checklist, Back-Out Plan,
Execution Summary, Approvals & Signatures.

Usage:
    python Scripts/generate_completion_docx.py

Output:
    Templates/MOP/MOP_Completion_Template.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Brand Colors ──
RESA_BLUE = RGBColor(0x01, 0x56, 0x87)
RESA_GREEN = RGBColor(0x5F, 0xA8, 0x44)
RESA_GRAY = RGBColor(0x68, 0x68, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)

# ── Paths ──
PACKAGE_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
NETA_LOGO = os.path.join(PACKAGE_ROOT, "assets", "neta-logo-90px.png")
RESA_LOGO = os.path.join(PACKAGE_ROOT, "assets", "resa-logo-90px.png")
OUTPUT_PATH = os.path.join(PACKAGE_ROOT, "artifacts", "MOP_Completion_Template.docx")


# ── Helper Functions (same pattern as cover/scope scripts) ──

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
            tcBorders.append(element)
        for key, val in attrs.items():
            element.set(qn(f"w:{key}"), str(val))


def add_formatted_run(paragraph, text, size=8, bold=False, color=BLACK, italic=False, caps=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    if caps:
        run.font.all_caps = True
    run.font.name = "Segoe UI"
    return run


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="exact"/>')
    trPr.append(trHeight)


def make_section_header(doc, text, alert=False):
    color_hex = "DC3545" if alert else "015687"
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color_hex)
    p = cell.paragraphs[0]
    add_formatted_run(p, text, size=8, bold=True, color=WHITE, caps=True)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for border_name in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": color_hex}})
    return table


def add_spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.line_spacing = Pt(pts)
    return p


def remove_all_borders(cell):
    for border_name in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{border_name: {"sz": 0, "val": "none", "color": "FFFFFF"}})


def add_branded_header(doc):
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(3.5)
    header_table.columns[2].width = Inches(2.0)

    cell_left = header_table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = cell_left.paragraphs[0]
    if os.path.exists(NETA_LOGO):
        p.add_run().add_picture(NETA_LOGO, height=Inches(0.5))
    else:
        add_formatted_run(p, "[NETA Logo]", size=7, color=RESA_GRAY, italic=True)

    header_table.cell(0, 1).paragraphs[0].text = ""

    cell_right = header_table.cell(0, 2)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(RESA_LOGO):
        p.add_run().add_picture(RESA_LOGO, height=Inches(0.5))
    else:
        add_formatted_run(p, "[RESA Logo]", size=7, color=RESA_GRAY, italic=True)

    for row in header_table.rows:
        for cell in row.cells:
            remove_all_borders(cell)

    # Blue line
    blue_line = doc.add_table(rows=1, cols=1)
    blue_line.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = blue_line.cell(0, 0)
    set_cell_shading(cell, "015687")
    set_row_height(blue_line.rows[0], 2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for b in ["top", "left", "bottom", "right"]:
        set_cell_border(cell, **{b: {"sz": 0, "val": "none", "color": "015687"}})

    # Tagline + Address
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.columns[0].width = Inches(4.0)
    info_table.columns[1].width = Inches(3.5)

    p = info_table.cell(0, 0).paragraphs[0]
    add_formatted_run(p, "Reliable and Safe. ", size=7.5, bold=True, color=RESA_BLUE, italic=True)
    add_formatted_run(p, "The Power of Us.", size=7.5, bold=True, color=RESA_GREEN, italic=True)

    p = info_table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, "2447 W 12th St. Suite 5 · Tempe, AZ 85142", size=7.5, color=RESA_GRAY)

    for row in info_table.rows:
        for cell in row.cells:
            remove_all_borders(cell)


def add_doc_info_bar(doc):
    bar = doc.add_table(rows=1, cols=2)
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar.columns[0].width = Inches(4.0)
    bar.columns[1].width = Inches(3.5)

    cell = bar.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    add_formatted_run(p, "Method of Procedure  ", size=9, bold=True, color=RESA_BLUE)
    add_formatted_run(p, "{{mop_id}} · v{{version}}", size=7, color=RESA_GRAY)

    cell = bar.cell(0, 1)
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, "{{project_name}}\n{{site_name}}", size=7, color=RESA_GRAY)

    for cell in bar.row_cells(0):
        set_cell_border(cell,
            top={"sz": 0, "val": "none", "color": "FFFFFF"},
            left={"sz": 0, "val": "none", "color": "FFFFFF"},
            right={"sz": 0, "val": "none", "color": "FFFFFF"},
            bottom={"sz": 4, "val": "single", "color": "DDDDDD"},
        )


def add_branded_footer(doc, page_text="Page {{page}} of {{total_pages}}"):
    add_spacer(doc, 4)
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.columns[0].width = Inches(1.5)
    footer_table.columns[1].width = Inches(6.0)

    green_cell = footer_table.cell(0, 0)
    set_cell_shading(green_cell, "5FA844")
    p = green_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "resapower.com", size=6, bold=True, color=WHITE)

    blue_cell = footer_table.cell(0, 1)
    set_cell_shading(blue_cell, "015687")
    p = blue_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_run(p, f"RESA Power Service  ·  NETA Accredited Company  ·  {page_text}", size=6, color=WHITE)

    for cell in [green_cell, blue_cell]:
        remove_all_borders(cell)


# ══════════════════════════════════════════════
#  MAIN DOCUMENT BUILD
# ══════════════════════════════════════════════

def build_document():
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(8)
    style.font.color.rgb = BLACK

    # ═══════════════════  HEADER  ═══════════════════
    add_branded_header(doc)
    add_doc_info_bar(doc)
    add_spacer(doc, 4)

    # ═══════════════════  POST-WORK VERIFICATION  ═══════════════════
    make_section_header(doc, "Post-Work Verification & Restoration Checklist")

    scope_completion = [
        "MV Switching complete",
        "LV Switching complete",
        "Relay Testing complete",
        "Protection Testing complete",
        "Commissioning complete",
        "LOTO removed / transferred",
        "CT Circuits restored",
        "PT Circuits restored",
        "Outage coordination restored",
        "Control Center notified",
        "Hot Work Permit closed",
        "Confined Space cleared",
    ]
    restoration = [
        "All scope items verified complete",
        "Equipment operating normally",
        "No unexpected alarms or faults",
        "All connections torqued to spec",
        "Covers and panels reinstalled",
        "No tools or materials left behind",
        "Work area cleaned and restored",
        "Barriers and signage removed",
        "Fire system returned to normal",
        "All personnel accounted for",
        "Documentation complete",
        "Customer walkthrough completed",
    ]

    num_rows = max(len(scope_completion), len(restoration))
    verify_table = doc.add_table(rows=num_rows + 1, cols=4)
    verify_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    verify_table.columns[0].width = Inches(0.35)
    verify_table.columns[1].width = Inches(3.40)
    verify_table.columns[2].width = Inches(0.35)
    verify_table.columns[3].width = Inches(3.40)

    # Column titles
    cell = verify_table.cell(0, 0)
    cell.merge(verify_table.cell(0, 1))
    p = cell.paragraphs[0]
    add_formatted_run(p, "SCOPE COMPLETION", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/></w:pBdr>'))

    cell = verify_table.cell(0, 2)
    cell.merge(verify_table.cell(0, 3))
    p = cell.paragraphs[0]
    add_formatted_run(p, "RESTORATION CHECKLIST", size=7, bold=True, color=RESA_BLUE)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="5FA844"/></w:pBdr>'))

    # Checklist rows — all blank checkboxes for post-work
    for i in range(num_rows):
        row = verify_table.rows[i + 1]
        if i < len(scope_completion):
            p = row.cells[0].paragraphs[0]
            add_formatted_run(p, "☐", size=9, color=RESA_BLUE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = row.cells[1].paragraphs[0]
            add_formatted_run(p, scope_completion[i], size=7.5, color=BLACK)

        if i < len(restoration):
            p = row.cells[2].paragraphs[0]
            add_formatted_run(p, "☐", size=9, color=RESA_BLUE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = row.cells[3].paragraphs[0]
            add_formatted_run(p, restoration[i], size=7.5, color=BLACK)

    for row in verify_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 0, "val": "none", "color": "FFFFFF"},
                left={"sz": 0, "val": "none", "color": "FFFFFF"},
                right={"sz": 0, "val": "none", "color": "FFFFFF"},
                bottom={"sz": 0, "val": "none", "color": "FFFFFF"},
            )

    # Outer border
    tbl = verify_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="015687"/>'
        f'</w:tblBorders>'
    ))

    # ═══════════════════  BACK-OUT PLAN  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "⚠ Back-Out Plan", alert=True)

    backout_table = doc.add_table(rows=7, cols=2)
    backout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    backout_table.columns[0].width = Inches(3.75)
    backout_table.columns[1].width = Inches(3.75)

    # Subtitles
    cell = backout_table.cell(0, 0)
    set_cell_shading(cell, "FFF3F3")
    p = cell.paragraphs[0]
    add_formatted_run(p, "TRIGGER CONDITIONS", size=7, bold=True, color=RGBColor(0xDC, 0x35, 0x45))
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="DC3545"/></w:pBdr>'))

    cell = backout_table.cell(0, 1)
    set_cell_shading(cell, "FFF3F3")
    p = cell.paragraphs[0]
    add_formatted_run(p, "BACK-OUT PROCEDURE", size=7, bold=True, color=RGBColor(0xDC, 0x35, 0x45))
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="DC3545"/></w:pBdr>'))

    triggers = [
        "Unexpected condition not covered by MOP",
        "Safety concern identified during work",
        "Unable to complete scope within window",
        "Customer-initiated abort",
        "Equipment damage or failure discovered",
    ]
    procedures = [
        "1. Stop all work — secure equipment",
        "2. Notify crew lead & customer PM",
        "3. Restore system to safe state",
        "4. Reinstall original configuration",
        "5. Complete restoration switching per MOP",
        "6. Document all actions & deviations",
    ]

    max_items = max(len(triggers), len(procedures))
    for i in range(max_items):
        row = backout_table.rows[i + 1]
        if i < len(triggers):
            p = row.cells[0].paragraphs[0]
            set_cell_shading(row.cells[0], "FFF3F3")
            add_formatted_run(p, f"• {triggers[i]}", size=7.5, color=BLACK)
        else:
            set_cell_shading(row.cells[0], "FFF3F3")

        if i < len(procedures):
            p = row.cells[1].paragraphs[0]
            set_cell_shading(row.cells[1], "FFF3F3")
            add_formatted_run(p, procedures[i], size=7.5, bold=True, color=BLACK)
        else:
            set_cell_shading(row.cells[1], "FFF3F3")

    # Border styling
    for row in backout_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 0, "val": "none", "color": "FFFFFF"},
                left={"sz": 0, "val": "none", "color": "FFFFFF"},
                right={"sz": 0, "val": "none", "color": "FFFFFF"},
                bottom={"sz": 0, "val": "none", "color": "FFFFFF"},
            )
    # Red outer border
    tbl = backout_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="DC3545"/>'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="DC3545"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="DC3545"/>'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="DC3545"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="DC3545"/>'
        f'</w:tblBorders>'
    ))

    # ═══════════════════  EXECUTION SUMMARY  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Execution Summary")

    exec_table = doc.add_table(rows=2, cols=4)
    exec_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    time_fields = [
        ("Planned Start", "{{planned_start}}"),
        ("Actual Start", ""),
        ("Planned End", "{{planned_end}}"),
        ("Actual End", ""),
    ]
    for i, (label, value) in enumerate(time_fields):
        # Label row
        cell = exec_table.cell(0, i)
        set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        add_formatted_run(p, label, size=6.5, bold=True, color=RESA_GRAY, caps=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Value row
        cell = exec_table.cell(1, i)
        p = cell.paragraphs[0]
        add_formatted_run(p, value, size=8, bold=True, color=RESA_BLUE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in exec_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 2, "val": "single", "color": "DDDDDD"},
                left={"sz": 2, "val": "single", "color": "DDDDDD"},
                right={"sz": 2, "val": "single", "color": "DDDDDD"},
                bottom={"sz": 2, "val": "single", "color": "DDDDDD"},
            )

    # Deviations
    add_spacer(doc, 2)
    dev_table = doc.add_table(rows=2, cols=1)
    dev_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = dev_table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    add_formatted_run(p, "DEVIATIONS FROM PLAN", size=6.5, bold=True, color=RESA_GRAY, caps=True)

    cell = dev_table.cell(1, 0)
    p = cell.paragraphs[0]
    add_formatted_run(p, "(Document any deviations from the original MOP, including scope changes, "
                        "timing changes, and unexpected conditions encountered.)", size=7, italic=True, color=RESA_GRAY)
    p.paragraph_format.space_after = Pt(24)

    for row in dev_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 2, "val": "single", "color": "DDDDDD"},
                left={"sz": 2, "val": "single", "color": "DDDDDD"},
                right={"sz": 2, "val": "single", "color": "DDDDDD"},
                bottom={"sz": 2, "val": "single", "color": "DDDDDD"},
            )

    # ═══════════════════  APPROVALS & SIGNATURES  ═══════════════════
    add_spacer(doc, 4)
    make_section_header(doc, "Approvals & Signatures")

    # Digital/Wet signature option
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "☐ ", size=9, color=RESA_BLUE)
    add_formatted_run(p, "Digital signatures accepted", size=7, color=RESA_GRAY)
    add_formatted_run(p, "    ☐ ", size=9, color=RESA_BLUE)
    add_formatted_run(p, "Wet signatures required", size=7, color=RESA_GRAY)

    add_spacer(doc, 2)

    # 4 signature blocks in 2×2 grid
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.columns[0].width = Inches(3.75)
    sig_table.columns[1].width = Inches(3.75)

    sig_blocks = [
        ("Prepared By", "RESA Power Service"),
        ("Reviewed By", "RESA Power Service"),
        ("Approved By", "Customer / Site Authority"),
        ("Customer Acceptance", "Customer Representative"),
    ]

    for i, (role, org) in enumerate(sig_blocks):
        row_idx = (i // 2) * 2
        col_idx = i % 2

        # Title row
        cell = sig_table.cell(row_idx, col_idx)
        set_cell_shading(cell, "F5F5F5")
        p = cell.paragraphs[0]
        add_formatted_run(p, role, size=7, bold=True, color=RESA_BLUE)
        add_formatted_run(p, f"  ({org})", size=6.5, color=RESA_GRAY)

        # Fields row
        cell = sig_table.cell(row_idx + 1, col_idx)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        add_formatted_run(p, "Name / Title: ", size=6.5, color=RESA_GRAY)
        add_formatted_run(p, "________________________________\n", size=7, color=RESA_GRAY)
        add_formatted_run(p, "Signature:       ", size=6.5, color=RESA_GRAY)
        add_formatted_run(p, "________________________________\n", size=7, color=RESA_GRAY)
        add_formatted_run(p, "Date:               ", size=6.5, color=RESA_GRAY)
        add_formatted_run(p, "________________________________", size=7, color=RESA_GRAY)
        p.paragraph_format.space_after = Pt(4)

    for row in sig_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": 2, "val": "single", "color": "DDDDDD"},
                left={"sz": 2, "val": "single", "color": "DDDDDD"},
                right={"sz": 2, "val": "single", "color": "DDDDDD"},
                bottom={"sz": 2, "val": "single", "color": "DDDDDD"},
            )

    # ═══════════════════  FOOTER  ═══════════════════
    add_branded_footer(doc, "Page {{page}} of {{total_pages}}")

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"✅ Word template saved: {OUTPUT_PATH}")
    print(f"   File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
